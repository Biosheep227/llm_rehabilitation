import cv2
import mediapipe as mp
import math
import pyttsx3
import time
import keyboard  # 用于检测按键输入
import os  # 用于文件路径操作
from datetime import datetime  # 用于生成时间戳
from docx import Document
from docx.shared import Inches

# 初始化MediaPipe和pyttsx3
mp_pose = mp.solutions.pose
pose = mp_pose.Pose()
mp_drawing = mp.solutions.drawing_utils
engine = pyttsx3.init()

# 语音播报函数
def speak(message):
    print(message)  # 让无声音设备也能看到问题
    engine.say(message)
    engine.runAndWait()

# 计算角度
def calculate_angle(a, b, c):
    ab = [b[0] - a[0], b[1] - a[1]]
    bc = [c[0] - b[0], c[1] - b[1]]
    dot_product = ab[0] * bc[0] + ab[1] * bc[1]
    cross_product = ab[0] * bc[1] - ab[1] * bc[0]
    angle = math.degrees(math.atan2(abs(cross_product), dot_product))
    return angle

# 询问用户，并使用语音播报问题
def ask_question(question):
    speak(question)
    print(question + "（按 'y' 表示是，按 'n' 表示否）")  # 让无声音设备也能看到问题
    while True:
        if keyboard.is_pressed('y'):
            return "是"
        elif keyboard.is_pressed('n'):
            return "否"

# 生成包含表格的 Word 报告
def generate_report_with_image(image_path, report_content, report_path):
    doc = Document()
    doc.add_heading('姿势检测报告', 0)
    doc.add_paragraph(f"报告生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # 创建表格，包含列标题：症状、描述
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '症状'
    hdr_cells[1].text = '描述'

    # 设置单元格宽度（调整列宽）
    for cell in hdr_cells:
        cell.width = Inches(3)  # 设置每个表格单元格的宽度
    
    # 根据用户的回答添加症状描述
    symptoms = report_content.split('；')  # 假设症状描述以“；”分隔
    for symptom in symptoms:
        row_cells = table.add_row().cells
        symptom_parts = symptom.split('：')  # 根据“：”分隔症状和描述
        if len(symptom_parts) == 2:
            row_cells[0].text = symptom_parts[0]  # 症状
            row_cells[1].text = symptom_parts[1]  # 描述

    # 添加图片
    doc.add_paragraph("照片：")
    doc.add_picture(image_path, width=Inches(3))
    
    # 保存报告
    doc.save(report_path)

# 获取桌面路径
desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")

# 初始化摄像头
cap = cv2.VideoCapture(0)
last_time = time.time()
hip_angle_history = []
leg_raised = False

while cap.isOpened():
    ret, frame = cap.read()
    if not ret:
        continue

    frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
    results = pose.process(frame_rgb)

    if results.pose_landmarks:
        landmarks = results.pose_landmarks.landmark

        left_shoulder = (landmarks[mp_pose.PoseLandmark.LEFT_SHOULDER].x, landmarks[mp_pose.PoseLandmark.LEFT_SHOULDER].y)
        left_hip = (landmarks[mp_pose.PoseLandmark.LEFT_HIP].x, landmarks[mp_pose.PoseLandmark.LEFT_HIP].y)
        left_knee = (landmarks[mp_pose.PoseLandmark.LEFT_KNEE].x, landmarks[mp_pose.PoseLandmark.LEFT_KNEE].y)
        left_ankle = (landmarks[mp_pose.PoseLandmark.LEFT_ANKLE].x, landmarks[mp_pose.PoseLandmark.LEFT_ANKLE].y)
        left_big_toe = (landmarks[mp_pose.PoseLandmark.LEFT_FOOT_INDEX].x, landmarks[mp_pose.PoseLandmark.LEFT_FOOT_INDEX].y)

        hip_angle = calculate_angle(left_shoulder, left_hip, left_knee)
        knee_angle = calculate_angle(left_hip, left_knee, left_ankle)
        ankle_angle = calculate_angle(left_knee, left_ankle, left_big_toe)

        hip_angle_complement = 180 - hip_angle
        knee_angle_complement = 180 - knee_angle
        ankle_angle_complement = 180 - ankle_angle

        # 在左上角显示角度信息
        cv2.putText(frame, f"hip: {hip_angle_complement:.2f}", (20, 40), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 0), 2)
        cv2.putText(frame, f"knee: {knee_angle_complement:.2f}", (20, 80), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 0), 2)
        cv2.putText(frame, f"ankle: {ankle_angle_complement:.2f}", (20, 120), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 0), 2)

        if (100 < ankle_angle_complement < 110) and (173 < knee_angle_complement < 180) and (175 < hip_angle_complement < 180):
            speak("请抬左腿，保持膝盖打直和脚掌勾起，抬至最高点保持三秒")
            leg_raised = True

        # 检测用户是否抬腿并检查膝关节
        if leg_raised:
            # 如果髋关节补角小于150度，开始每秒检测一次角度
            if hip_angle_complement < 150:
                # 获取当前时间
                current_time = time.time()
                
                # 每秒检测一次
                if current_time - last_time >= 1:  # 每1秒检测一次
                    last_time = current_time  # 更新最后检测时间
                
                    if knee_angle_complement < 170:
                        speak("膝关节未伸直")
                    else:
                        # 记录髋关节角度历史
                        if len(hip_angle_history) > 0:
                            hip_angle_diff = abs(hip_angle_complement - hip_angle_history[-1])
                        else:
                            hip_angle_diff = float('inf')  # 第一次检测时，不进行角度比较

                        # 记录当前髋关节角度
                        hip_angle_history.append(hip_angle_complement)
                        if len(hip_angle_history) > 3:  # 只保留最新的 2 次数据
                            hip_angle_history.pop(0)

                        # 判断条件：只有连续三次角度差 < 5°，才拍摄照片
                        if len(hip_angle_history) > 2 and hip_angle_diff < 5:
                            # 生成时间戳
                            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                            
                            # 照片文件路径
                            file_path = os.path.join(desktop_path, f"pose_{timestamp}.jpg")

                           # 在照片左上角标注膝关节和髋关节补角
                            cv2.putText(frame, f"Lift angle: {180-hip_angle_history[2]:.2f}", (20, 160),
                                        cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 0), 2)

                            # 保存照片
                            cv2.imwrite(file_path, frame)
                            speak("照片已拍摄并保存")

                           # 询问用户问题并记录
                            speak("请回答以下问题，若答案为是则按y，若答案为否则按n")

                            answer_1 = ask_question("是否麻木？")
                            if answer_1 == "是":
                                report_content = "神经：坐骨神经卡压、紧张；关节：髋关节囊僵紧；\
软组织：（梨状肌、股二头肌）水肿、卡压，血液循环差；代偿：骨盆、腰椎代偿；无力：核心、臀肌、腘绳肌无力；\
调整：松解、拉伸受限组织，训练肌肉离心能力，增强肌力、稳定性与协调性"#对应表格第一行
                            else:
                                answer_2 = ask_question("是否疼痛或不适？")
                                if answer_2 == "是":
                                    if hip_angle > 70:
                                        report_content = "神经：腰骶丛神经紧张；关节：腰椎小关节僵紧；\
软组织：腰方肌、腰大肌、臀肌紧张，血液循环差；代偿：骨盆代偿；无力：核心无力；\
调整：松解、拉伸受限组织，训练肌肉离心能力，增强肌力、稳定性与协调性"#对应表格第三行
                                    else :
                                        report_content = "神经：坐骨神经紧张；关节：骶髂关节、髋关节囊僵紧；\
软组织：后表链紧张，臀肌、腘绳肌、小腿三头肌紧张，血液循环差；代偿：腰椎代偿；无力：髂腰肌、后表链肌群紧张且无力；\
调整：松解、拉伸受限组织，训练肌肉离心能力，增强肌力、稳定性与协调性"#对应表格第二行
                                else:
                                    answer_3 = ask_question("是否牵扯？")
                                    if answer_3 == "是":
                                        if hip_angle > 70:
                                            report_content = "神经：腰骶丛神经紧张；关节：髋关节囊僵紧；\
软组织：腘绳肌紧张，血液循环差；代偿：骨盆代偿；无力：核心无力；\
调整：松解、拉伸受限组织，训练肌肉离心能力，增强肌力、稳定性与协调性"#对应表格第五行
                                        else :
                                            report_content = "神经：坐骨神经紧张；关节：髋关节囊僵紧；\
软组织： 臀肌、腘绳肌、小腿三头肌紧张，血液循环差；代偿：腰椎代偿；无力：髂腰肌、后表链肌群紧张且无力；\
调整：松解、拉伸受限组织，训练肌肉离心能力，增强肌力、稳定性与协调性"#对应表格第四行
                                    else:
                                        report_content = "未出现明显症状"

                            # 生成报告
                            report_path = os.path.join(desktop_path, f"report_{timestamp}.docx")
                            generate_report_with_image(file_path, report_content, report_path)

                            speak(f"报告已生成并保存至桌面")
                            cap.release()
                            cv2.destroyAllWindows()
                            break

    cv2.imshow("Pose Detection", frame)

    if cv2.waitKey(1) & 0xFF == ord('q'):
        break

cap.release()
cv2.destroyAllWindows()
