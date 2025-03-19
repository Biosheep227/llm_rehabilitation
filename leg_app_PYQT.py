import sys
import cv2
import mediapipe as mp
import math
import pyttsx3
import time
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches

from PyQt5.QtWidgets import (QApplication, QMainWindow, QLabel, QLineEdit, QPushButton,
                             QVBoxLayout, QWidget, QMessageBox)
from PyQt5.QtCore import QTimer, Qt
from PyQt5.QtGui import QImage, QPixmap

class PoseApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("姿势检测")
        self.subject_id = ""
        
        # UI控件：标签、输入框和按钮
        self.input_label = QLabel("请输入受试者编号:")
        self.subject_input = QLineEdit()
        self.start_button = QPushButton("开始检测")
        self.start_button.clicked.connect(self.start_detection)
        
        # 用于显示摄像头画面的标签
        self.video_label = QLabel()
        self.video_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        layout = QVBoxLayout()
        layout.addWidget(self.input_label)
        layout.addWidget(self.subject_input)
        layout.addWidget(self.start_button)
        layout.addWidget(self.video_label)
        
        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)
        
        # 初始化 mediapipe 和 pyttsx3
        self.mp_pose = mp.solutions.pose
        self.pose = self.mp_pose.Pose()
        self.mp_drawing = mp.solutions.drawing_utils
        self.engine = pyttsx3.init()
        
        # 摄像头和计时器
        self.cap = None
        self.timer = QTimer()
        self.timer.timeout.connect(self.update_frame)
        
        # 用于逻辑判断的变量
        self.last_time = time.time()
        self.hip_angle_history = []
        self.leg_raised = False

    def speak(self, message):
        # 语音播报并打印消息
        print(message)
        self.engine.say(message)
        self.engine.runAndWait()

    def calculate_angle(self, a, b, c):
        # 计算三点之间的夹角
        ab = [b[0] - a[0], b[1] - a[1]]
        bc = [c[0] - b[0], c[1] - b[1]]
        dot_product = ab[0] * bc[0] + ab[1] * bc[1]
        cross_product = ab[0] * bc[1] - ab[1] * bc[0]
        angle = math.degrees(math.atan2(abs(cross_product), dot_product))
        return angle

    def ask_question(self, question):
        # 使用消息对话框进行简单的“是/否”提问
        msg_box = QMessageBox(self)
        msg_box.setWindowTitle("提问")
        msg_box.setText(question + "\n点击 [是] 或 [否]")
        yes_button = msg_box.addButton("是", QMessageBox.ButtonRole.YesRole)
        no_button = msg_box.addButton("否", QMessageBox.ButtonRole.NoRole)
        msg_box.exec()
        if msg_box.clickedButton() == yes_button:
            return "是"
        else:
            return "否"

    def generate_report_with_image(self, image_path, report_content, report_path):
        doc = Document()
        doc.add_heading('姿势检测报告', 0)
        doc.add_paragraph(f"报告生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        # 创建表格
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '症状'
        hdr_cells[1].text = '描述'
        for cell in hdr_cells:
            cell.width = Inches(3)
        symptoms = report_content.split('；')
        for symptom in symptoms:
            row_cells = table.add_row().cells
            symptom_parts = symptom.split('：')
            if len(symptom_parts) == 2:
                row_cells[0].text = symptom_parts[0]
                row_cells[1].text = symptom_parts[1]
        doc.add_paragraph("照片：")
        doc.add_picture(image_path, width=Inches(3))
        doc.save(report_path)

    def start_detection(self):
        # 获取受试者编号，并启动摄像头检测
        self.subject_id = self.subject_input.text().strip()
        if not self.subject_id:
            QMessageBox.warning(self, "警告", "请输入受试者编号!")
            return
        # camera
        # self.cap = cv2.VideoCapture(1)
        self.cap = cv2.VideoCapture(0)
        if not self.cap.isOpened():
            QMessageBox.critical(self, "错误", "无法打开摄像头!")
            return
        self.timer.start(30)  # 每30毫秒更新一次画面
        self.start_button.setEnabled(False)
        self.subject_input.setEnabled(False)

    def update_frame(self):
        ret, frame = self.cap.read()
        if not ret:
            return
        # 转为RGB并处理姿势检测
        frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        results = self.pose.process(frame_rgb)
        if results.pose_landmarks:
            landmarks = results.pose_landmarks.landmark
            left_shoulder = (landmarks[self.mp_pose.PoseLandmark.LEFT_SHOULDER.value].x,
                             landmarks[self.mp_pose.PoseLandmark.LEFT_SHOULDER.value].y)
            left_hip = (landmarks[self.mp_pose.PoseLandmark.LEFT_HIP.value].x,
                        landmarks[self.mp_pose.PoseLandmark.LEFT_HIP.value].y)
            left_knee = (landmarks[self.mp_pose.PoseLandmark.LEFT_KNEE.value].x,
                         landmarks[self.mp_pose.PoseLandmark.LEFT_KNEE.value].y)
            left_ankle = (landmarks[self.mp_pose.PoseLandmark.LEFT_ANKLE.value].x,
                          landmarks[self.mp_pose.PoseLandmark.LEFT_ANKLE.value].y)
            left_big_toe = (landmarks[self.mp_pose.PoseLandmark.LEFT_FOOT_INDEX.value].x,
                            landmarks[self.mp_pose.PoseLandmark.LEFT_FOOT_INDEX.value].y)
            hip_angle = self.calculate_angle(left_shoulder, left_hip, left_knee)
            knee_angle = self.calculate_angle(left_hip, left_knee, left_ankle)
            ankle_angle = self.calculate_angle(left_knee, left_ankle, left_big_toe)
            hip_angle_complement = 180 - hip_angle
            knee_angle_complement = 180 - knee_angle
            ankle_angle_complement = 180 - ankle_angle

            # 在画面上显示角度信息
            cv2.putText(frame, f"hip: {hip_angle_complement:.2f}", (20, 40), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 0), 2)
            cv2.putText(frame, f"knee: {knee_angle_complement:.2f}", (20, 80), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 0), 2)
            cv2.putText(frame, f"ankle: {ankle_angle_complement:.2f}", (20, 120), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 0), 2)

            cv2.line(frame, (int(left_shoulder[0] * frame.shape[1]), int(left_shoulder[1] * frame.shape[0])), 
                    (int(left_hip[0] * frame.shape[1]), int(left_hip[1] * frame.shape[0])), (0, 255, 0), 2)  # 连接肩膀和髋
            cv2.line(frame, (int(left_hip[0] * frame.shape[1]), int(left_hip[1] * frame.shape[0])), 
                    (int(left_knee[0] * frame.shape[1]), int(left_knee[1] * frame.shape[0])), (0, 255, 0), 2)  # 连接髋和膝
            cv2.line(frame, (int(left_knee[0] * frame.shape[1]), int(left_knee[1] * frame.shape[0])), 
                    (int(left_ankle[0] * frame.shape[1]), int(left_ankle[1] * frame.shape[0])), (0, 255, 0), 2)  # 连接膝和踝
            cv2.line(frame, (int(left_ankle[0] * frame.shape[1]), int(left_ankle[1] * frame.shape[0])), 
                    (int(left_big_toe[0] * frame.shape[1]), int(left_big_toe[1] * frame.shape[0])), (0, 255, 0), 2)  # 连接踝和大拇指


            cv2.circle(frame, (int(left_shoulder[0] * frame.shape[1]), int(left_shoulder[1] * frame.shape[0])), 5, (0, 0, 255), -1)  # 左肩
            cv2.circle(frame, (int(left_hip[0] * frame.shape[1]), int(left_hip[1] * frame.shape[0])), 5, (0, 255, 255), -1)  # 左髋
            cv2.circle(frame, (int(left_knee[0] * frame.shape[1]), int(left_knee[1] * frame.shape[0])), 5, (255, 0, 0), -1)  # 左膝
            cv2.circle(frame, (int(left_ankle[0] * frame.shape[1]), int(left_ankle[1] * frame.shape[0])), 5, (255, 255, 0), -1)  # 左踝
            cv2.circle(frame, (int(left_big_toe[0] * frame.shape[1]), int(left_big_toe[1] * frame.shape[0])), 5, (255, 0, 255), -1)  # 左脚大拇指

            # 满足条件时，提示受试者抬腿
            if (100 < ankle_angle_complement < 110) and (173 < knee_angle_complement < 180) and (175 < hip_angle_complement < 180):
                self.speak("请抬左腿，保持膝盖打直和脚掌勾起，抬至最高点保持三秒")
                self.leg_raised = True

            if self.leg_raised:
                if hip_angle_complement < 150:
                    current_time = time.time()
                    if current_time - self.last_time >= 1:
                        self.last_time = current_time
                        if knee_angle_complement < 170:
                            self.speak("膝关节未伸直")
                        else:
                            if len(self.hip_angle_history) > 0:
                                hip_angle_diff = abs(hip_angle_complement - self.hip_angle_history[-1])
                            else:
                                hip_angle_diff = float('inf')
                            self.hip_angle_history.append(hip_angle_complement)
                            if len(self.hip_angle_history) > 3:
                                self.hip_angle_history.pop(0)
                            if len(self.hip_angle_history) > 2 and hip_angle_diff < 5:
                                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                                file_name = f"pose_{self.subject_id}_{timestamp}.jpg"
                                file_path = os.path.join(os.path.expanduser("~"), "Desktop", file_name)
                                cv2.putText(frame, f"Lift angle: {180-self.hip_angle_history[2]:.2f}", (20, 160),
                                            cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 0), 2)
                                cv2.imwrite(file_path, frame)
                                self.speak("照片已拍摄并保存")
                                
                                self.speak("请回答以下问题，若答案为是则点击[是]，否则点击[否]")
                                answer_1 = self.ask_question("是否麻木？")
                                if answer_1 == "是":
                                    report_content = (
                                        "神经：坐骨神经卡压、紧张；关节：髋关节囊僵紧；"
                                        "软组织：（梨状肌、股二头肌）水肿、卡压，血液循环差；"
                                        "代偿：骨盆、腰椎代偿；无力：核心、臀肌、腘绳肌无力；"
                                        "调整：松解、拉伸受限组织，训练肌肉离心能力，增强肌力、稳定性与协调性"
                                    )
                                else:
                                    answer_2 = self.ask_question("是否疼痛或不适？")
                                    if answer_2 == "是":
                                        if hip_angle > 70:
                                            report_content = (
                                                "神经：腰骶丛神经紧张；关节：腰椎小关节僵紧；"
                                                "软组织：腰方肌、腰大肌、臀肌紧张，血液循环差；"
                                                "代偿：骨盆代偿；无力：核心无力；"
                                                "调整：松解、拉伸受限组织，训练肌肉离心能力，增强肌力、稳定性与协调性"
                                            )
                                        else:
                                            report_content = (
                                                "神经：坐骨神经紧张；关节：骶髂关节、髋关节囊僵紧；"
                                                "软组织：后表链紧张，臀肌、腘绳肌、小腿三头肌紧张，血液循环差；"
                                                "代偿：腰椎代偿；无力：髂腰肌、后表链肌群紧张且无力；"
                                                "调整：松解、拉伸受限组织，训练肌肉离心能力，增强肌力、稳定性与协调性"
                                            )
                                    else:
                                        answer_3 = self.ask_question("是否牵扯？")
                                        if answer_3 == "是":
                                            if hip_angle > 70:
                                                report_content = (
                                                    "神经：腰骶丛神经紧张；关节：髋关节囊僵紧；"
                                                    "软组织：腘绳肌紧张，血液循环差；代偿：骨盆代偿；"
                                                    "无力：核心无力；调整：松解、拉伸受限组织，训练肌肉离心能力，增强肌力、稳定性与协调性"
                                                )
                                            else:
                                                report_content = (
                                                    "神经：坐骨神经紧张；关节：髋关节囊僵紧；"
                                                    "软组织：臀肌、腘绳肌、小腿三头肌紧张，血液循环差；"
                                                    "代偿：腰椎代偿；无力：髂腰肌、后表链肌群紧张且无力；"
                                                    "调整：松解、拉伸受限组织，训练肌肉离心能力，增强肌力、稳定性与协调性"
                                                )
                                        else:
                                            report_content = "未出现明显症状"
                                report_name = f"report_{self.subject_id}_{timestamp}.docx"
                                report_path = os.path.join(os.path.expanduser("~"), "Desktop", report_name)
                                self.generate_report_with_image(file_path, report_content, report_path)
                                self.speak(f"报告已生成并保存至桌面")
                                self.timer.stop()
                                self.cap.release()
                                cv2.destroyAllWindows()
        
        # 将 frame 转为 QImage 显示到 QLabel 上
        frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        h, w, ch = frame_rgb.shape
        bytes_per_line = ch * w
        q_img = QImage(frame_rgb.data, w, h, bytes_per_line, QImage.Format.Format_RGB888)
        self.video_label.setPixmap(QPixmap.fromImage(q_img))
        
    def closeEvent(self, event):
        self.timer.stop()
        if self.cap is not None:
            self.cap.release()
        cv2.destroyAllWindows()
        event.accept()

if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = PoseApp()
    window.show()
    sys.exit(app.exec())