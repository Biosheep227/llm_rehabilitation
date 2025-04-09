import cv2
import mediapipe as mp
import pyttsx3
import math
import threading
import time
import os
import sys
from docx import Document
from docx.shared import Inches

# 转换系数（如需用于其他距离计算，此处保留，仅作参考）
PIXELS_PER_CM = 37.0

# 初始化 MediaPipe Pose 模型
mp_pose = mp.solutions.pose
pose = mp_pose.Pose()

# 初始化语音引擎
engine = pyttsx3.init()

def calculate_angle(a, b, c):
    """
    计算由点 a, b, c 构成的角度，其中 b 为顶点。
    a, b, c 均为 [x, y] 格式的坐标。
    """
    ab = [a[0] - b[0], a[1] - b[1]]
    bc = [c[0] - b[0], c[1] - b[1]]
    dot_product = ab[0] * bc[0] + ab[1] * bc[1]
    magnitude_ab = math.sqrt(ab[0]**2 + ab[1]**2)
    magnitude_bc = math.sqrt(bc[0]**2 + bc[1]**2)
    if magnitude_ab * magnitude_bc == 0:
        return 0
    cos_angle = dot_product / (magnitude_ab * magnitude_bc)
    cos_angle = max(min(cos_angle, 1.0), -1.0)
    angle = math.degrees(math.acos(cos_angle))
    return angle

def speak(text):
    """语音播报函数，采用多线程避免阻塞主循环"""
    threading.Thread(target=lambda: (engine.say(text), engine.runAndWait())).start()

def draw_point(point, frame_width, frame_height):
    """将归一化坐标转换为像素坐标"""
    return (int(point[0] * frame_width), int(point[1] * frame_height))

def euclidean_distance(p1, p2):
    """计算两点之间的欧氏距离"""
    return math.hypot(p1[0] - p2[0], p1[1] - p2[1])

# 语音提示相关设置
last_spoken = {"arms": 0, "knee": 0, "ankle": 0, "first_posture": 0}
cooldown = 2  # 单位：秒

has_spoken_first = False
is_first_message_spoken = False

# 当播报“请将双手向前伸展”后开始记录手部信息（不记录轨迹）
start_tracking = False

cap = cv2.VideoCapture(0)

# 用于保存最近 3 秒内髋关节角度数据，格式为 (时间戳, hip_angle)
hip_angle_history = []

# 记录髋关节稳定开始的时间
stability_start_time = None

# 防止重复触发拍照生成报告
capture_triggered = False

while cap.isOpened():
    ret, frame = cap.read()
    if not ret:
        break

    # 保存一份原始帧（未绘制额外信息），用于最终报告
    frame_original = frame.copy()

    frame_height, frame_width, _ = frame.shape
    rgb_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
    results = pose.process(rgb_frame)
    current_time = time.time()

    if results.pose_landmarks:
        landmarks = results.pose_landmarks.landmark

        # 获取左侧关键点
        left_shoulder = [landmarks[mp_pose.PoseLandmark.LEFT_SHOULDER].x,
                         landmarks[mp_pose.PoseLandmark.LEFT_SHOULDER].y]
        left_hip = [landmarks[mp_pose.PoseLandmark.LEFT_HIP].x,
                    landmarks[mp_pose.PoseLandmark.LEFT_HIP].y]
        left_knee = [landmarks[mp_pose.PoseLandmark.LEFT_KNEE].x,
                     landmarks[mp_pose.PoseLandmark.LEFT_KNEE].y]
        left_ankle = [landmarks[mp_pose.PoseLandmark.LEFT_ANKLE].x,
                      landmarks[mp_pose.PoseLandmark.LEFT_ANKLE].y]
        left_foot_index = [landmarks[mp_pose.PoseLandmark.LEFT_FOOT_INDEX].x,
                           landmarks[mp_pose.PoseLandmark.LEFT_FOOT_INDEX].y]
        left_elbow = [landmarks[mp_pose.PoseLandmark.LEFT_ELBOW].x,
                      landmarks[mp_pose.PoseLandmark.LEFT_ELBOW].y]
        left_wrist = [landmarks[mp_pose.PoseLandmark.LEFT_WRIST].x,
                      landmarks[mp_pose.PoseLandmark.LEFT_WRIST].y]
        left_index_landmark = landmarks[mp_pose.PoseLandmark.LEFT_INDEX]

        # 在实时画面中绘制骨架（不显示手部轨迹）
        cv2.line(frame, draw_point(left_shoulder, frame_width, frame_height),
                 draw_point(left_hip, frame_width, frame_height), (0, 255, 0), 2)
        cv2.line(frame, draw_point(left_hip, frame_width, frame_height),
                 draw_point(left_knee, frame_width, frame_height), (0, 255, 0), 2)
        cv2.line(frame, draw_point(left_knee, frame_width, frame_height),
                 draw_point(left_ankle, frame_width, frame_height), (0, 255, 0), 2)
        cv2.line(frame, draw_point(left_ankle, frame_width, frame_height),
                 draw_point(left_foot_index, frame_width, frame_height), (0, 255, 0), 2)

        # 计算各关节角度
        knee_angle = calculate_angle(left_hip, left_knee, left_ankle)
        ankle_angle = calculate_angle(left_knee, left_ankle, left_foot_index)
        # 对脚踝角度减10度
        adjusted_ankle_angle = ankle_angle - 10
        hip_angle = calculate_angle(left_shoulder, left_hip, left_knee)  # 左侧髋关节角度

        cv2.putText(frame, f"Knee: {knee_angle:.1f}", (10, 30),
                    cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 0), 2)
        # 显示时仍显示原始脚踝角度（可根据需要改为 adjusted_ankle_angle）
        cv2.putText(frame, f"Ankle: {ankle_angle:.1f}", (10, 60),
                    cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 0), 2)
        cv2.putText(frame, f"Hip: {hip_angle:.1f}", (10, 90),
                    cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 0), 2)

        # 定义腿部姿势要求，脚踝条件使用调整后的角度
        baseline_leg = (173 <= knee_angle <= 180) and (100 <= adjusted_ankle_angle <= 120)

        # 首次符合条件时播报“请坐直并将双手打平”
        if baseline_leg and not has_spoken_first and not is_first_message_spoken:
            if current_time - last_spoken["first_posture"] > cooldown:
                speak("请坐直并将双手打平")
                has_spoken_first = True
                is_first_message_spoken = True
                last_spoken["first_posture"] = current_time

        # 后续提示
        if is_first_message_spoken:
            if baseline_leg and (80 <= hip_angle <= 100):
                if current_time - last_spoken["arms"] > cooldown:
                    speak("请将双手向前伸展")
                    last_spoken["arms"] = current_time
                    if not start_tracking:
                        start_tracking = True

            left_shoulder_elbow_angle = calculate_angle(left_shoulder, left_elbow, [left_elbow[0], left_elbow[1] + 1])
            cv2.putText(frame, f"Left Elbow-Shr Angle: {left_shoulder_elbow_angle:.1f}°", (10, 120),
                        cv2.FONT_HERSHEY_SIMPLEX, 1, (255, 255, 0), 2)
            if not (173 <= knee_angle <= 180):
                if current_time - last_spoken["knee"] > cooldown:
                    speak("膝关节未伸直")
                    last_spoken["knee"] = current_time
            # 脚踝判断条件采用调整后的角度：原条件 (100 <= ankle_angle <= 120) 改为 (100 <= adjusted_ankle_angle <= 120)
            if not (90 <= adjusted_ankle_angle <= 110):
                if current_time - last_spoken["ankle"] > cooldown:
                    speak("脚踝未勾起")
                    last_spoken["ankle"] = current_time

        # 记录当前髋关节角度及时间（保留最近 3 秒数据）
        hip_angle_history.append((current_time, hip_angle))
        hip_angle_history = [(t, a) for (t, a) in hip_angle_history if current_time - t <= 3]

        # 仅当 start_tracking 为 True 且当前髋关节角度 < 80°时记录稳定时间
        if start_tracking and hip_angle < 80:
            if hip_angle_history:
                angles = [a for (_, a) in hip_angle_history]
                if max(angles) - min(angles) < 5:
                    if stability_start_time is None:
                        stability_start_time = hip_angle_history[0][0]
                    stable_duration = current_time - stability_start_time
                    cv2.putText(frame, f"Stability: {stable_duration:.1f} s", (10, 150),
                                cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 255), 2)
                    # 当稳定时间达到 5 秒且未触发拍照时生成报告
                    if stable_duration >= 5 and not capture_triggered:
                        capture_triggered = True

                        # 使用原始帧生成最终报告图（未绘制额外信息）
                        captured_image_original = frame_original.copy()

                        # 在 captured_image_original 上绘制骨架与角度信息
                        cv2.line(captured_image_original, draw_point(left_shoulder, frame_width, frame_height),
                                 draw_point(left_hip, frame_width, frame_height), (0, 255, 0), 2)
                        cv2.line(captured_image_original, draw_point(left_hip, frame_width, frame_height),
                                 draw_point(left_knee, frame_width, frame_height), (0, 255, 0), 2)
                        cv2.line(captured_image_original, draw_point(left_knee, frame_width, frame_height),
                                 draw_point(left_ankle, frame_width, frame_height), (0, 255, 0), 2)
                        cv2.line(captured_image_original, draw_point(left_ankle, frame_width, frame_height),
                                 draw_point(left_foot_index, frame_width, frame_height), (0, 255, 0), 2)
                        cv2.putText(captured_image_original, f"Knee: {knee_angle:.1f}", (10, 30),
                                    cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 0), 2)
                        cv2.putText(captured_image_original, f"Ankle: {ankle_angle:.1f}", (10, 60),
                                    cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 0), 2)
                        cv2.putText(captured_image_original, f"Hip: {hip_angle:.1f}", (10, 90),
                                    cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 0), 2)

                        # 计算各关键点坐标
                        pt_index = draw_point([left_index_landmark.x, left_index_landmark.y], frame_width, frame_height)  # 指尖
                        pt_wrist = draw_point(left_wrist, frame_width, frame_height)  # 掌根
                        # 指根：手腕与指尖连线中点（0.5处）
                        finger_root = (int(pt_wrist[0] + 0.5 * (pt_index[0] - pt_wrist[0])),
                                       int(pt_wrist[1] + 0.5 * (pt_index[1] - pt_wrist[1])))
                        pt_toe = draw_point(left_foot_index, frame_width, frame_height)  # 脚尖（通常对应大脚趾尖）
                        pt_ankle = draw_point(left_ankle, frame_width, frame_height)      # 脚踝

                        # 在最终报告图像上显示各关键点坐标（显示在屏幕左上方）
                        cv2.putText(captured_image_original,
                                    f"Finger tip: {pt_index}",
                                    (10, 70), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 0, 255), 2)
                        cv2.putText(captured_image_original,
                                    f"Finger root: {finger_root}",
                                    (10, 110), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 0, 255), 2)
                        cv2.putText(captured_image_original,
                                    f"Palm root: {pt_wrist}",
                                    (10, 150), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 0, 255), 2)
                        cv2.putText(captured_image_original,
                                    f"Ankle: {pt_ankle}",
                                    (10, 190), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 0, 255), 2)
                        cv2.putText(captured_image_original,
                                    f"Toe: {pt_toe}",
                                    (10, 230), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 0, 255), 2)

                        # 判断接触情况：先用距离判断，再用横坐标判断
                        threshold_distance = 30  # 像素阈值
                        dist_wrist = euclidean_distance(pt_wrist, pt_toe)
                        dist_finger_root = euclidean_distance(finger_root, pt_toe)
                        dist_index = euclidean_distance(pt_index, pt_toe)
                        if dist_wrist < threshold_distance:
                            judgement = "掌根刚好碰到脚尖"
                        elif dist_finger_root < threshold_distance:
                            judgement = "指根刚好碰到脚尖"
                        elif dist_index < threshold_distance:
                            judgement = "指尖刚好碰到脚尖"
                        else:
                            if pt_toe[0] < pt_index[0]:
                                judgement = "指尖碰不到脚尖"
                            elif pt_wrist[0] < pt_toe[0]:
                                judgement = "掌根已超过脚尖"
                            else:
                                judgement = "请重新测试"

                        final_image_filename = "final_trajectory.jpg"
                        cv2.imwrite(final_image_filename, captured_image_original)

                        # 生成 Word 报告
                        doc = Document()
                        doc.add_heading('姿态检测报告', 0)
                        doc.add_paragraph(f'判断结果：{judgement}')
                        doc.add_paragraph('拍摄照片如下（显示指尖、指根、掌根、脚踝和脚尖坐标）：')
                        doc.add_picture(final_image_filename, width=Inches(4))
                        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
                        report_filename = os.path.join(desktop_path, f'Pose_Report_{int(current_time)}.docx')
                        doc.save(report_filename)

                        cv2.putText(frame, "Report Generated!", (50, frame_height - 50),
                                    cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 0, 255), 2)
                        speak("报告已生成并保存至桌面")
                        cv2.imshow("Pose Detection", frame)
                        cv2.waitKey(500)
                        time.sleep(2)
                        break
                else:
                    stability_start_time = None
            else:
                stability_start_time = None
        else:
            stability_start_time = None

    cv2.imshow("Pose Detection", frame)
    if cv2.waitKey(1) & 0xFF == ord('q'):
        break

cap.release()
cv2.destroyAllWindows()
sys.exit(0)
